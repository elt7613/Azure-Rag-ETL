"""Local parsers: pdfplumber (PDF), python-docx (DOCX), openpyxl (XLSX)."""
from __future__ import annotations

import io
import re

from rag.models import Block, BlockType, ParsedDocument

_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(\S.*)$")
_HEADER_HINT = "|"
# pdfplumber's raw PDF font-code fallback for an undecoded bullet glyph.
_CID_RE = re.compile(r"\(cid:\d+\)\s*")
# Real pipe-delimited metadata lines carry 3-4 fields; a single stray "|" can
# turn up in ordinary prose or a table fragment, so require at least two to
# treat a line as the document's metadata header.
_MIN_HEADER_PIPES = 2


def _level_of(number: str) -> int:
    return number.count(".") + 1


def _is_header_line(text: str) -> bool:
    return text.count(_HEADER_HINT) >= _MIN_HEADER_PIPES


# A running header/footer repeats on most pages of a document. Digits are
# masked before comparison so "... Page 1" and "... Page 2" collapse to the
# same signature; without that, a footer carrying a page number looks unique on
# every page and survives.
_DIGITS_RE = re.compile(r"\d+")
# Furniture lives in the top and bottom band of the page. A sentence repeated
# mid-page is far more likely to be genuine repeated policy text.
_EDGE_BAND_RATIO = 0.12
# Below this many pages, "repeats on most pages" is not evidence of anything --
# a two-page document can legitimately open both pages with the same phrase.
_MIN_PAGES_FOR_FURNITURE = 2


def _furniture_signature(text: str) -> str:
    return _DIGITS_RE.sub("#", " ".join(text.split())).lower()


def _detect_page_furniture(
    pages: list[list[tuple[float, str]]], page_heights: list[float]
) -> set[str]:
    """Signatures of lines that are running headers/footers, not content.

    Detected structurally rather than by pattern: a line is furniture when the
    same digit-masked text appears in the top or bottom band of at least half
    the pages. That catches a footer like "Contoso Ltd. — Internal
    Use Only Page 2" without hard-coding the company name, and it will catch
    whatever the next corpus uses instead.

    Left in place, these lines are pure noise in three places at once: they
    become their own chunks, they dilute the embedding of every chunk they land
    in, and they burn context tokens in the answer prompt.
    """
    if len(pages) < _MIN_PAGES_FOR_FURNITURE:
        return set()

    seen: dict[str, set[int]] = {}
    for page_index, (entries, height) in enumerate(zip(pages, page_heights)):
        if not height:
            continue
        band = height * _EDGE_BAND_RATIO
        for top, text in entries:
            if top > band and top < height - band:
                continue
            seen.setdefault(_furniture_signature(text), set()).add(page_index)

    threshold = max(2, (len(pages) + 1) // 2)
    return {sig for sig, page_indices in seen.items() if len(page_indices) >= threshold}


class LocalParser:
    async def parse(self, data: bytes, doc_id: str) -> ParsedDocument:
        suffix = doc_id.rsplit(".", 1)[-1].lower()
        if suffix == "pdf":
            return self._parse_pdf(data, doc_id)
        if suffix == "docx":
            return self._parse_docx(data, doc_id)
        if suffix == "xlsx":
            return self._parse_xlsx(data, doc_id)
        raise ValueError(f"unsupported format: {suffix}")

    # ---------------- PDF ----------------
    def _parse_pdf(self, data: bytes, doc_id: str) -> ParsedDocument:
        import pdfplumber

        blocks: list[Block] = []
        header = ""
        # Two passes: collect every page's laid-out entries first so running
        # headers/footers can be identified by their repetition across pages,
        # then emit blocks with those lines removed. A single streaming pass
        # cannot do this -- on page 1 a footer is indistinguishable from a
        # sentence.
        page_entries: list[list[tuple[float, str, object]]] = []
        page_text_lines: list[list[tuple[float, str]]] = []
        page_heights: list[float] = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page_no, page in enumerate(pdf.pages, start=1):
                tables = page.find_tables()
                table_entries: list[tuple[float, Block]] = []
                for t in tables:
                    raw = t.extract() or []
                    rows = [[(c or "").replace("\n", " ").strip() for c in r] for r in raw]
                    if rows:
                        table_entries.append((t.bbox[1], Block(type=BlockType.TABLE, text="",
                                                                page=page_no, rows=rows)))

                # Exclude table regions from text extraction: pdfplumber's
                # extract_text() re-emits every table cell as plain text too,
                # which would otherwise (a) duplicate table content as
                # paragraphs and (b) let numeric table rows like
                # "0 - 2 years 15 days ..." false-match the heading regex.
                if tables:
                    def _outside_tables(obj, _tables=tables):
                        return not any(
                            t.bbox[0] <= obj.get("x0", 0) <= t.bbox[2]
                            and t.bbox[1] <= obj.get("top", 0) <= t.bbox[3]
                            for t in _tables
                        )
                    text_lines = page.filter(_outside_tables).extract_text_lines()
                else:
                    text_lines = page.extract_text_lines()

                line_entries = [(ln["top"], ln["text"].strip())
                                for ln in text_lines if ln["text"].strip()]

                # Interleave tables with text in true reading order (vertical
                # position on the page), rather than dumping every table
                # before any text. Emitting tables first meant a table's
                # heading hadn't been seen yet, so build_section_tree filed
                # every table under the synthetic "Preamble" section instead
                # of the heading it actually sits under.
                merged: list[tuple[float, str, object]] = sorted(
                    [(top, "table", block) for top, block in table_entries]
                    + [(top, "text", text) for top, text in line_entries],
                    key=lambda entry: entry[0],
                )
                page_entries.append(merged)
                page_text_lines.append(line_entries)
                page_heights.append(float(page.height or 0))

        furniture = _detect_page_furniture(page_text_lines, page_heights)

        for page_no, merged in enumerate(page_entries, start=1):
            i = 0
            while i < len(merged):
                _, kind, payload = merged[i]
                if kind == "table":
                    blocks.append(payload)
                    i += 1
                    continue
                line = payload
                if _furniture_signature(line) in furniture:
                    i += 1
                    continue
                if not header and _is_header_line(line):
                    header = line
                    # The metadata line sometimes wraps onto the next
                    # physical line as a single trailing word (e.g.
                    # "Policy Owner: VP, People" / "Operations"). Fold
                    # that continuation back in rather than truncating.
                    if i + 1 < len(merged) and merged[i + 1][1] == "text":
                        nxt = merged[i + 1][2]
                        if (len(nxt.split()) <= 2 and _HEADER_HINT not in nxt
                                and not _HEADING_RE.match(nxt)
                                and not any(ch.isdigit() for ch in nxt)):
                            header = f"{header} {nxt}"
                            i += 1
                    i += 1
                    continue
                blocks.append(self._line_to_block(line, page_no))
                i += 1
        return ParsedDocument(doc_id=doc_id, file_format="pdf",
                              blocks=blocks, raw_header=header)

    # ---------------- DOCX ----------------
    def _parse_docx(self, data: bytes, doc_id: str) -> ParsedDocument:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = docx.Document(io.BytesIO(data))
        blocks: list[Block] = []
        header = ""

        # Walk the body's XML children rather than `document.paragraphs` then
        # `document.tables`. Those two collections are separate and each is in
        # document order only within itself, so consuming them in sequence puts
        # every table after every paragraph -- which files all of them under
        # whichever heading happens to come last. In TravelPolicy.docx that put
        # the flight-class, hotel-cap and per-diem tables under "10 Contact"
        # instead of the sections that actually introduce them, so a question
        # about hotel caps could retrieve the right table with the wrong
        # citation. The PDF path already interleaves by page position; this is
        # the DOCX equivalent.
        for child in document.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                para = Paragraph(child, document)
                text = para.text.strip()
                if not text:
                    continue
                if not header and _is_header_line(text):
                    header = text
                    continue
                style = para.style.name if para.style is not None else ""
                style = (style or "").lower()
                if style.startswith("heading"):
                    digits = "".join(ch for ch in style if ch.isdigit())
                    blocks.append(Block(type=BlockType.HEADING, text=text,
                                        page=1, level=int(digits or 1)))
                else:
                    blocks.append(self._line_to_block(text, 1))
            elif tag == "tbl":
                table = Table(child, document)
                rows = [[c.text.strip() for c in r.cells] for r in table.rows]
                if rows:
                    blocks.append(
                        Block(type=BlockType.TABLE, text="", page=1, rows=rows)
                    )

        return ParsedDocument(doc_id=doc_id, file_format="docx",
                              blocks=blocks, raw_header=header)

    # ---------------- XLSX ----------------
    def _parse_xlsx(self, data: bytes, doc_id: str) -> ParsedDocument:
        import openpyxl

        # data_only=True reads cached computed values instead of formula text.
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        blocks: list[Block] = []
        header = ""
        for sheet in wb:
            blocks.append(Block(type=BlockType.HEADING, text=sheet.title,
                                page=1, level=1))
            group: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if not any(cells):
                    # Blank row: boundary between one logical table/paragraph
                    # run and the next.
                    if group:
                        blocks.extend(self._xlsx_group_to_blocks(group))
                        group = []
                    continue
                joined = " ".join(c for c in cells if c)
                if _is_header_line(joined):
                    # The "Sales Operations | Effective: ... | Version ..."
                    # metadata line repeats on every sheet. Strip it out of
                    # every sheet's data (not just the first) so it never
                    # leaks in as a bogus table row; only the first
                    # occurrence in the whole workbook becomes raw_header.
                    if not header:
                        header = joined
                    continue
                group.append(cells)
            if group:
                blocks.extend(self._xlsx_group_to_blocks(group))
        return ParsedDocument(doc_id=doc_id, file_format="xlsx",
                              blocks=blocks, raw_header=header)

    @staticmethod
    def _xlsx_group_to_blocks(group: list[list[str]]) -> list[Block]:
        """Turn one blank-row-delimited run of sheet rows into blocks.

        A run is only rendered as a TABLE (with rows[0] as the markdown
        header) when its first row is genuinely a column-header row -- i.e.
        it populates at least as many cells as any data row in the run.
        Sheet titles and label/value blocks (like the "Stacking Example")
        only populate their first column and would otherwise get rendered
        as the table header, burying the real column names. Those become
        plain paragraphs instead, one per row.
        """
        def populated(cells: list[str]) -> int:
            return sum(1 for c in cells if c)

        if (len(group) >= 2 and populated(group[0]) >= 2
                and populated(group[0]) >= max(populated(r) for r in group)):
            return [Block(type=BlockType.TABLE, text="", page=1, rows=group)]

        blocks: list[Block] = []
        for row in group:
            text = " ".join(c for c in row if c)
            if text:
                blocks.append(Block(type=BlockType.PARAGRAPH, text=text, page=1))
        return blocks

    # ---------------- shared ----------------
    def _line_to_block(self, line: str, page: int) -> Block:
        # pdfplumber sometimes fails to map a bullet glyph to a real
        # character and emits the raw PDF font code instead, e.g.
        # "(cid:127) List prices increased ...". Left in place this both
        # pollutes what's shown to users in citations and pollutes the
        # embedding vector. A leading "(cid:N)" is always a bullet marker
        # in this corpus, so strip it and classify the line as a LIST item;
        # any other occurrence is normalised away too.
        is_bullet = bool(_CID_RE.match(line))
        cleaned = _CID_RE.sub("", line).strip()
        if is_bullet:
            return Block(type=BlockType.LIST, text=cleaned, page=page)
        match = _HEADING_RE.match(cleaned)
        if match and len(match.group(2)) < 80:
            number, title = match.groups()
            return Block(type=BlockType.HEADING, text=f"{number} {title}",
                         page=page, level=_level_of(number))
        return Block(type=BlockType.PARAGRAPH, text=cleaned, page=page)
